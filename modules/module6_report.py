# =========================
# MODULE 6 : REPORT GENERATOR (STREAMLIT READY)
# =========================

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
from io import BytesIO
import re
import tempfile
import os

# =========================
# CONSTANTS
# =========================
ID_MONTHS = {
    1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni",
    7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober",
    11: "November", 12: "Desember"
}

# =========================
# DATE PARSER
# =========================
def parse_date_flexible(date_str: str):
    if not date_str:
        return None

    s = str(date_str).strip()

    month_map = {
        "Januari":"January","Februari":"February","Maret":"March","April":"April",
        "Mei":"May","Juni":"June","Juli":"July","Agustus":"August",
        "September":"September","Oktober":"October","November":"November","Desember":"December"
    }
    for indo, eng in month_map.items():
        s = s.replace(indo, eng)

    fmts = [
        "%d %B %Y", "%d.%m.%Y", "%d-%m-%Y",
        "%Y-%m-%d", "%d/%m/%Y", "%d %b %Y", "%m/%d/%Y"
    ]
    for f in fmts:
        try:
            return datetime.strptime(s, f)
        except:
            pass

    m = re.search(r"(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})", s)
    if m:
        d, mn, y = m.groups()
        if len(y) == 2:
            y = "20" + y
        try:
            return datetime.strptime(f"{d}-{mn}-{y}", "%d-%m-%Y")
        except:
            pass

    m2 = re.search(r"(\d{4})[./-](\d{1,2})[./-](\d{1,2})", s)
    if m2:
        y, mn, d = m2.groups()
        try:
            return datetime.strptime(f"{d}-{mn}-{y}", "%d-%m-%Y")
        except:
            pass

    return None


def format_date_id(dt: datetime):
    if not dt:
        return ""
    return f"{dt.day:02d} {ID_MONTHS.get(dt.month,'')} {dt.year}"


def format_date_en(dt: datetime):
    if not dt:
        return ""
    return dt.strftime("%B %d, %Y")

# =========================
# STYLE UTILITIES
# =========================
def style_paragraph(p, size=12, bold=False, italic=False, align="left"):
    if not p.runs:
        p.add_run("")

    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except:
            pass

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    borders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)

    tblPr.append(borders)

# =========================
# SECTION BUILDERS (ASLI)
# =========================
def build_title(doc, row):
    dt = parse_date_flexible(row.get("Tanggal Koordinat",""))
    t_str = format_date_id(dt) if dt else row.get("Tanggal Koordinat","")

    ka = row.get("Koordinat Awal","")
    kb = row.get("Koordinat Akhir","")

    p = doc.add_paragraph()
    p.add_run("Meteorological Reports").bold = True
    p.add_run("\nCoordinate: From ").bold = True
    p.add_run(f"{ka} ")
    p.add_run("To ").bold = True
    p.add_run(f"{kb}\n")
    p.add_run(f"for {t_str}")

    style_paragraph(p, bold=True, align="center")
    doc.add_paragraph("")


def build_interval_table(doc, intervals, tz="WIB):
    headers = [
        "DATE",f"LOCAL TIME ({tz})","WEATHER",
        "WIND (Knot)","CURRENT (cm/s)",
        "WAVE (meter)","BEAUFORT SCALE"
    ]

    table = doc.add_table(rows=1, cols=7)
    set_table_border(table)

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        style_paragraph(table.rows[0].cells[i].paragraphs[0], bold=True, align="center")

    for j in range(4):
        data = intervals[j] if j < len(intervals) else {}
        row = table.add_row().cells
        values = [
            data.get("DATE",""),
            data.get("LOCAL TIME",""),
            data.get("WEATHER",""),
            data.get("WIND",""),
            data.get("CURRENT",""),
            data.get("WAVE",""),
            data.get("BEAUFORT",""),
        ]
        for i, v in enumerate(values):
            row[i].text = str(v)
            style_paragraph(row[i].paragraphs[0], align="center")

    doc.add_paragraph("")


def build_notes_primary(doc):
    p = doc.add_paragraph()
    p.add_run("Note:\n").bold = True
    p.add_run("The direction of current is toward.\nThe direction of wind is from.")
    style_paragraph(p, size=11, italic=True)
    doc.add_paragraph("")


def build_wave_category_table(doc):
    data = [
        ("Smooth","0.10 – 0.50 m"),
        ("Slight","0.50 – 1.25 m"),
        ("Moderate","1.25 – 2.50 m"),
        ("Rough","2.50 – 4.00 m"),
        ("Very Rough","4.00 – 6.00 m"),
        ("High","6.00 – 9.00 m"),
        ("Very High","9.00 – 14.00 m"),
    ]

    t = doc.add_table(rows=1, cols=2)
    set_table_border(t)

    for label,val in data:
        cells = t.add_row().cells
        cells[0].text = label
        cells[1].text = val
        style_paragraph(cells[0].paragraphs[0], size=11, align="center")
        style_paragraph(cells[1].paragraphs[0], size=11, align="center")

    doc.add_paragraph("")


def build_satellite_image_table(doc, tanggal_str):
    dt = parse_date_flexible(tanggal_str)
    tanggal_fmt = format_date_id(dt) if dt else tanggal_str

    table = doc.add_table(rows=2, cols=2)
    set_table_border(table)

    hdr = table.rows[0].cells[0]
    hdr.merge(table.rows[0].cells[1])

    p = hdr.paragraphs[0]
    p.add_run(f"Weather Satellite Image on {tanggal_fmt} at ______")
    style_paragraph(p, bold=True, align="center")

    table.rows[1].cells[0].paragraphs[0].add_run("[Insert Satellite Image Here]")
    table.rows[1].cells[1].paragraphs[0].add_run("[Insert Legend Here]")

    style_paragraph(table.rows[1].cells[0].paragraphs[0], italic=True, align="center")
    style_paragraph(table.rows[1].cells[1].paragraphs[0], italic=True, align="center")

    doc.add_paragraph("")

# =========================
# MAIN ENTRY (STREAMLIT)
# =========================
def generate_final_docx_streamlit(module1_rows, module5_rows, template_path):
    """
    Menghasilkan file DOCX siap download di Streamlit
    """

    doc = Document(template_path)

    # === COVER ===
    first = module1_rows[0]
    ka = first.get("Koordinat Awal", "")
    kb = first.get("Koordinat Akhir", "")

    dt_awal = parse_date_flexible(module1_rows[0].get("Tanggal Koordinat", ""))
    dt_akhir = parse_date_flexible(module1_rows[-1].get("Tanggal Koordinat", ""))

    if dt_awal and dt_akhir:
        periode = f"{dt_awal.strftime('%B %d')} - {dt_akhir.strftime('%d, %Y')}"
    else:
        periode = ""

    replacements = {
        "$nama_perusahaan": first.get("Nama Perusahaan", ""),
        "$alamat_perusahaan": first.get("Alamat Perusahaan", ""),
        "$no_surat": first.get("Nomor Surat", ""),
        "$LIST_KOORDINAT": f"From {ka} to {kb}\nfor {periode}",
        "$tanggal_hari_ini": datetime.now().strftime("%d %B %Y"),
    }

    for p in doc.paragraphs:
        for k, v in replacements.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
                style_paragraph(p)

    # === ISI ===
    for idx, row in enumerate(module1_rows):

        # Validasi keamanan sebelum akses module5_rows
        if not module5_rows:
            continue

        if idx >= len(module5_rows):
            continue

        module5_item = module5_rows[idx]

        if module5_item is None:
            continue

        if "intervals" not in module5_item:
            continue

        build_title(doc, row)

        intervals = module5_item["intervals"]
        tz = module5_item.get("tz", "WIB")
        build_interval_table(doc, intervals, tz)
        build_notes_primary(doc)
        build_wave_category_table(doc)
        build_satellite_image_table(
            doc,
            row.get("Tanggal Koordinat", "")
        )

        if idx < len(module1_rows) - 1:
            doc.add_page_break()

    # === SAVE TO MEMORY ===
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
